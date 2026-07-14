#!/usr/bin/env python3
"""Generate sqlBotAi technical solution document."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_doc_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(80, 80, 80)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    set_doc_font(doc)
    doc.sections[0].top_margin = Cm(2.5)
    doc.sections[0].bottom_margin = Cm(2.5)
    doc.sections[0].left_margin = Cm(2.8)
    doc.sections[0].right_margin = Cm(2.8)

    today = date.today().strftime("%Y年%m月%d日")

    add_title(doc, "AI 数据资产平台（sqlBotAi）技术方案")
    add_subtitle(doc, f"版本：V1.0    日期：{today}")
    doc.add_paragraph()

    add_heading(doc, "1. 文档说明", 1)
    add_table(
        doc,
        ["项目", "说明"],
        [
            ["文档名称", "AI 数据资产平台（sqlBotAi）技术方案"],
            ["系统名称", "AI 数据资产平台 / sqlBotAi"],
            ["版本号", "1.0.0"],
            ["编写依据", "当前工程代码、配置、部署脚本及 Wiki 知识库"],
            ["适用范围", "产品、研发、数据资产、运维及相关干系人"],
        ],
    )

    add_heading(doc, "2. 项目概述", 1)
    add_heading(doc, "2.1 建设背景", 2)
    add_para(
        doc,
        "企业在数据治理与数据资产运营过程中，积累了大量指标口径、表资产、血缘关系、"
        "业务规则与 SQL 实现，但分散在文档、脚本与业务人员经验中，查询成本高、口径不一致、"
        "难以支撑业务人员的自助分析。",
    )
    add_para(
        doc,
        "sqlBotAi 旨在构建一套面向企业数据资产的智能问答与分析平台，将本地 Wiki 知识库、"
        "向量检索（RAG）、大语言模型（LLM）与受控业务查数能力结合，实现“问口径能答、问数据能查”。",
    )

    add_heading(doc, "2.2 建设目标", 2)
    add_bullets(
        doc,
        [
            "统一承载数据资产知识：指标、表、概念、血缘、领域参考文档。",
            "提供知识问答能力：支持概念解释、口径说明、规则查询。",
            "提供业务查数能力：自然语言问题自动路由、生成受控 SQL、查询 MySQL 并输出分析结论。",
            "提供知识图谱可视化：基于 Wiki 链接与元数据构建可浏览的关系网络。",
            "支持标准化部署：一键打包、上传、安装、启动，适配阿里云 ECS 生产环境。",
        ],
    )

    add_heading(doc, "2.3 建设范围（当前版本）", 2)
    add_table(
        doc,
        ["范围项", "当前状态", "说明"],
        [
            ["Wiki 知识库", "已建设", "Markdown 文档，覆盖线下营销域进攻专案 MVP 等"],
            ["知识问答（RAG）", "已实现", "DashScope Embedding + 关键词混合检索"],
            ["业务查数", "已实现", "query-business-data Skill，线下营销域试点"],
            ["知识图谱", "已实现", "基于 Wiki 自动构建节点与边"],
            ["概念/实体/资产 CRUD", "已实现", "演示与元数据管理能力"],
            ["图表看板", "部分实现", "基于种子数据的统计图；业务查数图表脚本待接入"],
            ["用户认证授权", "未实现", "当前所有接口公开访问"],
            ["上传文档入 RAG", "未实现", "上传文件仅元数据管理，未进入向量索引"],
        ],
    )

    add_heading(doc, "3. 总体技术架构", 1)
    add_heading(doc, "3.1 架构分层", 2)
    add_table(
        doc,
        ["层次", "技术选型", "职责"],
        [
            ["展现层", "Thymeleaf + Bootstrap 5", "Wiki 浏览、知识问答、知识图谱、资产管理页面"],
            ["应用层", "Spring Boot 3.2 + Java 17", "路由、业务编排、API 服务"],
            ["智能层", "LangChain4j + DeepSeek + DashScope", "问答生成、SQL 生成、结果总结、向量嵌入"],
            ["知识层", "本地 Markdown Wiki + RAG 索引", "指标/表/概念/领域参考知识供给"],
            ["执行层", "Python query_mysql.py", "受控 SQL 执行、结果落盘"],
            ["数据层", "H2（元数据）+ MySQL ADS（业务数据）", "平台元数据与业务明细查询"],
            ["部署层", "Nginx + systemd + Shell", "反向代理、进程守护、自动化部署"],
        ],
    )

    add_heading(doc, "3.2 系统架构图（逻辑）", 2)
    add_para(doc, "用户 → Nginx(80) → Spring Boot(8080) → 问答路由")
    add_bullets(
        doc,
        [
            "概念/口径问题 → LocalRagService（Wiki 向量 + 关键词）→ DeepSeek 生成答案",
            "业务数据问题 → QueryBusinessDataService → Wiki 上下文 → SQL 生成 → MySQL 执行 → 结果总结",
            "无命中问题 → DeepSeek 通用回退",
            "图谱/目录 → WikiCatalogService / WikiKnowledgeGraphService → H2 图数据",
        ],
    )

    add_heading(doc, "3.3 技术栈清单", 2)
    add_table(
        doc,
        ["类别", "组件", "版本/说明"],
        [
            ["后端框架", "Spring Boot", "3.2.0"],
            ["语言", "Java", "17"],
            ["模板引擎", "Thymeleaf", "HTML 服务端渲染"],
            ["ORM", "Spring Data JPA", "H2 持久化"],
            ["LLM 接入", "LangChain4j OpenAiChatModel", "1.17.0，对接 DeepSeek"],
            ["Embedding", "LangChain4j QwenEmbeddingModel", "DashScope text-embedding-v2"],
            ["文档解析", "PDFBox / Apache POI", "上传文件解析（预留）"],
            ["查数执行", "Python 3 + query_mysql.py", "SSH/本地 MySQL 客户端"],
            ["构建工具", "Maven", "打包为可执行 JAR"],
        ],
    )

    add_heading(doc, "4. 功能模块设计", 1)

    modules = [
        ("4.1 首页与导航", "/", "系统入口，导航至 Wiki、知识问答、知识图谱、图表等模块。"),
        ("4.2 Wiki 知识库", "/wiki", "浏览 Markdown 知识目录树，查看指标、表资产、概念、领域参考文档。"),
        ("4.3 知识问答", "/knowledge-base", "核心问答入口，支持文档上传、RAG 问答、业务查数、来源标注。"),
        ("4.4 知识图谱", "/knowledge-graph", "基于 Wiki 链接与 frontmatter 构建 ECharts 关系图，支持重建。"),
        ("4.5 概念管理", "/concepts", "业务/技术概念 CRUD 与关键词检索。"),
        ("4.6 实体管理", "/entities", "表、API、文件等实体对象管理。"),
        ("4.7 数据资产管理", "/data-assets", "数据资产登记与维护。"),
        ("4.8 图表面板", "/charts", "基于种子数据的概念分布、实体趋势、资产统计图。"),
    ]
    for title, path, desc in modules:
        add_heading(doc, title, 2)
        add_para(doc, f"访问路径：{path}")
        add_para(doc, desc)

    add_heading(doc, "4.9 核心 API 一览", 2)
    add_table(
        doc,
        ["模块", "方法", "路径", "说明"],
        [
            ["知识问答", "POST", "/knowledge-base/api/chat", "统一问答入口"],
            ["知识问答", "GET", "/knowledge-base/api/documents", "已上传文档列表"],
            ["知识问答", "POST", "/knowledge-base/api/upload", "上传知识文档"],
            ["知识问答", "GET", "/knowledge-base/api/rag/stats", "RAG 索引统计"],
            ["知识问答", "POST", "/knowledge-base/api/rag/rebuild", "重建 RAG 索引"],
            ["Wiki", "GET", "/wiki/api/catalog", "Wiki 目录树"],
            ["Wiki", "GET", "/wiki/api/page", "Wiki 页面内容"],
            ["知识图谱", "GET", "/knowledge-graph/api/graph", "全量图谱数据"],
            ["知识图谱", "POST", "/knowledge-graph/api/rebuild", "从 Wiki 重建图谱"],
        ],
    )

    add_heading(doc, "5. 知识问答核心设计", 1)
    add_heading(doc, "5.1 问答路由策略", 2)
    add_para(doc, "KnowledgeBaseService.answerQuestion() 按以下优先级路由：")
    add_bullets(
        doc,
        [
            "第一优先级：业务查数（QueryIntentService 识别为数据查询）→ QueryBusinessDataService",
            "第二优先级：RAG 问答（rag.enabled=true）→ LocalRagService",
            "第三优先级：Wiki 关键词命中 → WikiSearchService 摘录回答",
            "第四优先级：DeepSeek 通用问答回退",
        ],
    )

    add_heading(doc, "5.2 意图识别（业务查数）", 2)
    add_para(
        doc,
        "QueryIntentService 通过正则识别“需要真实数据”的问题，关键词包括：多少、金额、数量、"
        "占比、达成率、最新一期、各期、趋势、Top N、排名、大区、省区、门店、奖励、进攻专案、轮动等。",
    )
    add_para(doc, "命中后进入 query-business-data Skill，不进入纯概念 RAG 路径。")

    add_heading(doc, "5.3 RAG 检索架构", 2)
    add_bullets(
        doc,
        [
            "知识源：classpath:wiki/**/*.md，启动时由 WikiChunkingService 分块（600 字，重叠 80）。",
            "向量索引：DashScope Embedding，缓存至 data/rag-cache；失败时回退 TF-IDF。",
            "检索策略：关键词检索 + 向量 Top-K 混合，去重合并后送入 LLM。",
            "生成策略：DeepSeekService.chatWithWikiKnowledge，温度 0.3，强制仅依据上下文回答。",
            "上下文上限：rag.max-context-length = 4000 字符。",
        ],
    )

    add_heading(doc, "5.4 Wiki 关键词检索增强", 2)
    add_bullets(
        doc,
        [
            "评分规则：标题命中 +5，路径 +3，正文 +1。",
            "业务同义词：轮动专案→进攻专案，奖励情况→最终奖励。",
            "领域强制注入：线下营销域问题自动注入领域参考、指标资产、奖励表、目标达成表页面。",
            "高相关文档（score≥5）注入全文，否则注入摘要片段。",
        ],
    )

    add_heading(doc, "6. 业务查数 Skill 设计", 1)
    add_heading(doc, "6.1 处理流程", 2)
    steps = [
        "接收用户自然语言问题。",
        "构建 Wiki 上下文：强制领域页 + 扩展关键词检索 + RAG 向量块。",
        "加载 Skill 提示词：SKILL.md、mysql-dialect.md、mvp-questions.md。",
        "调用 DeepSeek 生成单条 MySQL 8.4 只读 SQL（SELECT / WITH ... SELECT）。",
        "通过 query_mysql.py 执行 SQL，输出 result.json / result.csv。",
        "调用 DeepSeek 基于 Wiki 口径与查询结果生成中文业务分析。",
        "返回 ChatAnswerDTO：答案、来源、引用页、SQL、列名、数据行、截断标记。",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"{i}. {step}")

    add_heading(doc, "6.2 SQL 生成约束", 2)
    add_bullets(
        doc,
        [
            "仅允许 SELECT 或 WITH ... SELECT，禁止 DDL/DML/注释/多语句。",
            "表名使用 ADS/DIM/CDM 全限定名。",
            "轮动专案与进攻专案视为同一项目。",
            "奖励查询使用 final_reward，禁止使用 estimate_reward。",
            "省区字段 lev3_name，大区字段 lev2_name，禁止臆造 province_name。",
            "区域维度来自目标达成表，奖励来自最终奖励表，关联前需收敛粒度防膨胀。",
            "最新一期通过 MAX(data_date) 解析。",
        ],
    )

    add_heading(doc, "6.3 SQL 执行沙箱（query_mysql.py）", 2)
    add_table(
        doc,
        ["控制项", "策略"],
        [
            ["语句类型", "仅 SELECT / WITH ... SELECT"],
            ["危险关键字", "拦截 DROP/DELETE/UPDATE/INSERT/ALTER 等"],
            ["系统库", "禁止 information_schema 以外系统 schema 写操作"],
            ["函数限制", "拦截 sleep、load_file 等危险函数"],
            ["行数限制", "最多返回 500 行，超出标记 truncated"],
            ["超时控制", "MySQL MAX_EXECUTION_TIME = 30 秒"],
            ["连接方式", "开发：SSH 到远程服务器执行 mysql；生产：本地 mysql 客户端"],
        ],
    )

    add_heading(doc, "6.4 试点场景（线下营销域）", 2)
    add_table(
        doc,
        ["场景", "示例问题", "核心表/指标"],
        [
            ["大区最新奖励", "最新一期轮动专案各个大区的最终奖励金额", "final_reward + 目标达成表"],
            ["各期 Top10", "各期轮动专案奖励 Top10 连锁系统和单体店", "final_reward + terminal_filings_type"],
            ["省区达成率", "最新一期每个省区目标达成率", "目标达成表 lev3_name"],
            ["大区下省区奖励", "中一区下各省区最新一期奖励情况", "final_reward + lev2_name 过滤"],
        ],
    )

    add_heading(doc, "7. 知识图谱设计", 1)
    add_bullets(
        doc,
        [
            "数据源：Wiki Markdown 中的 [[双链]]、frontmatter concept_refs / source_tables 等。",
            "构建服务：WikiKnowledgeGraphService，启动时若图为空则自动重建。",
            "存储：H2 表 knowledge_graph_nodes / knowledge_graph_edges。",
            "展示：前端 ECharts 力导向图，支持按 Wiki 页面局部展开。",
            "接口：GET /knowledge-graph/api/graph、POST /api/rebuild。",
        ],
    )

    add_heading(doc, "8. 数据架构", 1)
    add_heading(doc, "8.1 平台元数据（H2）", 2)
    add_table(
        doc,
        ["实体", "表名", "用途"],
        [
            ["Concept", "concepts", "概念定义"],
            ["EntityItem", "entities", "实体对象"],
            ["DataAsset", "data_assets", "数据资产登记"],
            ["KnowledgeDocument", "knowledge_documents", "上传文档元数据"],
            ["KnowledgeGraphNode", "knowledge_graph_nodes", "图谱节点"],
            ["KnowledgeGraphEdge", "knowledge_graph_edges", "图谱边"],
        ],
    )

    add_heading(doc, "8.2 文件与缓存存储", 2)
    add_table(
        doc,
        ["路径", "内容"],
        [
            ["src/main/resources/wiki/", "Markdown 知识库（打包进 JAR）"],
            ["src/main/resources/knowledge-base/", "用户上传文档"],
            ["data/rag-cache/", "Embedding 向量缓存"],
            ["data/query-results/{sessionId}/", "业务查数 SQL 与结果"],
            ["/opt/sqlbotai/data/sqldb（生产）", "H2 文件数据库"],
            ["/opt/sqlbotai/config/env（生产）", "API Key 等环境变量"],
            ["/opt/sqlbotai/config/mysql.cnf（生产）", "MySQL 客户端凭据"],
        ],
    )

    add_heading(doc, "8.3 业务数据源（MySQL ADS）", 2)
    add_para(
        doc,
        "业务查数连接企业 ADS 层 MySQL 数据库。当前试点核心表包括："
        "ads_bi_attack_exemplary_case_final_reward_df（最终奖励）、"
        "ads_bi_attack_exemplary_case_target_achieve_df（目标达成）、"
        "ads_bi_attack_exemplary_case_base_data_df（基础数据）等。",
    )

    add_heading(doc, "9. 部署方案", 1)
    add_heading(doc, "9.1 部署拓扑", 2)
    add_bullets(
        doc,
        [
            "服务器：阿里云 ECS（示例 120.48.17.175）",
            "反向代理：Nginx 监听 80 端口，转发至 127.0.0.1:8080",
            "应用进程：systemd 服务 sqlbotai，用户 sqlbotai，Spring profile=prod",
            "应用目录：/opt/sqlbotai/{app,data,logs,config}",
            "业务数据库：同机或内网 MySQL，生产环境 connection-mode=local",
        ],
    )

    add_heading(doc, "9.2 部署流程", 2)
    add_table(
        doc,
        ["步骤", "脚本", "说明"],
        [
            ["1. 配置", "deploy/deploy.env", "填写 SERVER_HOST、SSH 凭据、API Key"],
            ["2. 打包", "deploy/package-artifacts.sh", "mvn package，生成 sqlbotai-deploy.tar.gz"],
            ["3. 上传部署", "deploy/deploy.sh", "SCP 上传并在服务器执行 install-on-server.sh"],
            ["4. 首次安装", "deploy/install-server.sh", "安装 Java17、Nginx、Python、systemd"],
            ["5. 升级安装", "deploy/install-on-server.sh", "替换 JAR、更新配置、重启服务"],
        ],
    )

    add_heading(doc, "9.3 运行时目录结构", 2)
    add_table(
        doc,
        ["目录", "说明"],
        [
            ["/opt/sqlbotai/app/", "sqlBotAi.jar 及 exploded 目录"],
            ["/opt/sqlbotai/data/", "H2 数据库、RAG 缓存、查数结果"],
            ["/opt/sqlbotai/logs/", "应用日志 sqlbotai.log"],
            ["/opt/sqlbotai/config/", "env、mysql.cnf 等敏感配置"],
        ],
    )

    add_heading(doc, "10. 配置说明", 1)
    add_table(
        doc,
        ["配置项", "配置键", "说明"],
        [
            ["服务端口", "server.port", "默认 8080，生产绑定 127.0.0.1"],
            ["DeepSeek", "deepseek.*", "LLM API Key、模型 deepseek-chat"],
            ["DashScope", "dashscope.*", "Embedding API、text-embedding-v2"],
            ["Wiki 检索", "wiki.*", "搜索模式、最小分数、最大结果数"],
            ["RAG", "rag.*", "分块、Top-K、向量存储、上下文长度"],
            ["业务查数", "query-business-data.*", "SSH/本地连接、行数上限、超时"],
        ],
    )

    add_heading(doc, "11. 安全设计", 1)
    add_heading(doc, "11.1 已落实措施", 2)
    add_bullets(
        doc,
        [
            "SQL 执行沙箱：只读查询、关键字黑名单、行数与超时限制。",
            "生产环境应用仅监听 localhost，由 Nginx 对外暴露。",
            "独立系统用户 sqlbotai 运行应用，配置文件 chmod 600。",
            "生产环境关闭 H2 Console。",
            "LLM 提示词要求不暴露数据库密码与内部配置。",
        ],
    )

    add_heading(doc, "11.2 待完善项", 2)
    add_bullets(
        doc,
        [
            "缺少用户认证与权限控制（Spring Security 未接入）。",
            "上传文件未做病毒扫描与访问鉴权。",
            "LLM 生成 SQL 仍依赖 Python 沙箱作为最终防线，需持续审计。",
            "API Key 以环境变量/文件方式管理，建议接入密钥托管服务。",
        ],
    )

    add_heading(doc, "12. 关键类与模块映射", 1)
    add_table(
        doc,
        ["类名", "包路径", "职责"],
        [
            ["KnowledgeBaseService", "service", "问答总路由"],
            ["QueryBusinessDataService", "service.query", "业务查数编排"],
            ["QueryIntentService", "service.query", "查数意图识别"],
            ["MySqlQueryExecutor", "service.query", "Python 查数执行器"],
            ["DeepSeekService", "service", "LLM 调用封装"],
            ["LocalRagService", "service.rag", "RAG 检索与生成"],
            ["WikiSearchService", "service", "Wiki 关键词检索"],
            ["WikiKnowledgeGraphService", "service", "知识图谱构建"],
            ["SkillPromptService", "service.query", "Skill 提示词加载"],
        ],
    )

    add_heading(doc, "13. 已知限制与后续规划", 1)
    add_table(
        doc,
        ["项", "现状", "规划建议"],
        [
            ["图表输出", "render_chart.py 已具备，Java 流程未调用", "接入查数结果自动出图"],
            ["上传文档 RAG", "仅 Wiki 进入向量索引", "扩展上传文档入索引"],
            ["用户体系", "无登录鉴权", "接入 SSO / Spring Security"],
            ["多域扩展", "线下营销域试点", "按领域参考文档扩展更多业务域"],
            ["语义层", "指标资产 formula 预埋", "升级为可执行语义指标层"],
            ["WebSocket/Neo4j", "依赖已引入未使用", "评估后接入或移除"],
        ],
    )

    add_heading(doc, "14. 附录：核心源码路径", 1)
    add_table(
        doc,
        ["类别", "路径"],
        [
            ["应用入口", "src/main/java/com/sqlbot/SqlBotAiApplication.java"],
            ["问答控制器", "src/main/java/com/sqlbot/controller/KnowledgeBaseController.java"],
            ["业务查数", "src/main/java/com/sqlbot/service/query/QueryBusinessDataService.java"],
            ["Skill 定义", "src/main/resources/skills/query-business-data/SKILL.md"],
            ["MySQL 脚本", "src/main/resources/skills/query-business-data/scripts/query_mysql.py"],
            ["Wiki 知识库", "src/main/resources/wiki/pages/"],
            ["生产配置", "src/main/resources/application-prod.yml"],
            ["部署脚本", "deploy/deploy.sh"],
        ],
    )

    doc.add_paragraph()
    add_para(doc, "—— 文档结束 ——", bold=True)

    return doc


def main() -> None:
    output = Path(__file__).resolve().parent / "sqlBotAi-技术方案.docx"
    doc = build_document()
    doc.save(output)
    print(f"已生成: {output}")


if __name__ == "__main__":
    main()
